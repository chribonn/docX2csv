from docx import Document

docName = "Test.docx"

document = Document(docName)

for para in document.paragraphs:
    print (para.text)
    print (para.style.name)
    print (para.style.type)
    print (para.contains_page_break)

    
# for para in document.paragraphs:
#     for run in para.runs:
#         print (run.text)
#         print (run.style)
#         print (run.contains_page_break)

    



# for paragraph in document.paragraphs:
#    print (paragraph.name, paragraph.text)






def list_sections_in_word_document(file_path):
    """
    Reads in a word document and lists all the sections in the document.

    Args:
    file_path (str): The file path of the word document.

    Returns:
    list: A list of section names in the document.
    """
    sections = []
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            sections.append(paragraph.text)
    return sections

